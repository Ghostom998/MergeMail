import docx, csv, os.path, bz2
import pandas as pd

# TODO ecept for "EMAIL" & "ATTACHMENT" we want to create variable <<SWAP>> words dynamically

# Get text from word document. This text will hold placeholder text such as <<FIRST>> for first names, etc.
def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

# TODO method to parse information from CSV to retrieve the text to fill the place holder. i.e. 'john' => <<FIRST>>
def getValues(filename):
    return pd.read_csv(filename)

# Method build new text based on previous
def BuildText(Text, dic):
    for key, item in dic.items():
        Text = Text.replace("<<" + key + ">>", item)
    return Text

""" for testing
def WriteNewDoc(Text, DocName):
    doc = docx.Document()
    doc.add_paragraph(Text)
    doc.save(DocName) """

"""         

with open(filename) as csvfile:
        reader = csv.DictReader(csvfile)
        fieldnames = reader.fieldnamesConstFields, DynFields = {}, {}

        # Generate Dictionaries
        while fieldnames:  # While there are arguments left to parse...
            NewField = [] 
            for row in reader:
                if fieldnames[0] != 'EMAIL' or 'ATTACHMENT':
                    DynFields[fieldnames[0]] = NewField.append(row[fieldnames[0]]) 
                else:
                    ConstFields[fieldnames[0]] = NewField.append(row[fieldnames[0]]) 
            fieldnames = fieldnames[1:]  # Reduce the argument list by copying it starting from index 1

    return ConstFields, DynFields """